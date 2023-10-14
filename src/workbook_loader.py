import os
from urllib.parse import urlparse
import requests
from openpyxl import load_workbook, Workbook
from docx import Document
from doc2docx import convert

from config import DOWNLOADED_SCHEDULES_DIR, USE_SCHEDULES_FROM_URLS, SCHEDULES_URLS


def download_file(url: str, destination_path: str) -> None:
    """Download a file from a given URL and save it to the specified destination path."""
    with requests.get(url, verify=False) as response:  # Handle SSL warnings properly
        response.raise_for_status()  # Ensure we got a valid response
        with open(destination_path, 'wb') as file:
            file.write(response.content)


def load_excel_from_path(path: str) -> Workbook:
    """Load an Excel workbook from a given path."""
    return load_workbook(path)


def convert_doc_to_docx_if_needed(path: str) -> str:
    """Convert a .doc file to .docx format if needed and return the .docx path."""
    if path.endswith('.doc'):
        docx_path = os.path.splitext(path)[0] + '.docx'
        convert(path, docx_path)
        while not os.path.exists(docx_path):
            pass
        os.remove(path)
        return docx_path
    return path


def load_word_as_excel(path: str) -> Workbook:
    """Load a Word (.doc or .docx) document and convert its content into an Excel workbook."""
    path = convert_doc_to_docx_if_needed(path)
    doc = Document(path)

    wb = Workbook()
    ws = wb.active

    for i, paragraph in enumerate(doc.paragraphs):
        cells = [sentence.strip() for sentence in paragraph.text.split('.') if sentence]
        for j, cell in enumerate(cells):
            ws.cell(row=i + 1, column=j + 1, value=cell)

    row_offset = len(doc.paragraphs)
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                ws.cell(row=row_offset + i + 1, column=j + 1, value=cell.text)

    return wb


def load_workbook_from_local_file(path: str):
    """Load a workbook from a local file. Supports Excel (.xlsx) and Word (.doc, .docx) files."""
    file_extension = os.path.splitext(path)[-1].lower()

    if file_extension == '.xlsx':
        return load_excel_from_path(path)
    elif file_extension in ['.doc', '.docx']:
        return load_word_as_excel(path)
    else:
        raise ValueError(f'Unsupported file extension: {file_extension}')


def load_workbook_from_url(url: str):
    """Download a workbook from a URL and load it. Supports Excel (.xlsx) and Word (.doc, .docx) files."""
    parsed_url = urlparse(url)
    filename = os.path.basename(parsed_url.path)

    local_path = os.path.join(DOWNLOADED_SCHEDULES_DIR, filename)
    if not os.path.exists(local_path):
        os.makedirs(DOWNLOADED_SCHEDULES_DIR, exist_ok=True)
        download_file(url, local_path)

    return load_workbook_from_local_file(local_path)


def get_workbook_sources():
    """Return a list of sources (URLs or local file paths) to load workbooks from."""
    allowed_extensions = ['.doc', '.docx', '.xlsx']
    if USE_SCHEDULES_FROM_URLS:
        return SCHEDULES_URLS
    else:
        return [os.path.join(DOWNLOADED_SCHEDULES_DIR, filename)
                for filename in os.listdir(DOWNLOADED_SCHEDULES_DIR)
                if os.path.splitext(filename)[-1].lower() in allowed_extensions and not filename.startswith('~$')]


def load_workbook_from_source(source):
    """Load a workbook from a source (either URL or local path) based on the configuration."""
    if USE_SCHEDULES_FROM_URLS:
        return load_workbook_from_url(source)
    else:
        return load_workbook_from_local_file(source)
